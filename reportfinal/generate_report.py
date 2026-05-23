"""
DevOps Project Report Generator – Nimma Yatri
Produces a fully formatted Word document + PDF via win32com.
"""

import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# PATHS
# ─────────────────────────────────────────────
BASE   = os.path.dirname(os.path.abspath(__file__))
PICS   = os.path.join(BASE, "pics")
DOCX   = os.path.join(BASE, "DevOps_Project_Report_NimmaYatri.docx")
PDF    = os.path.join(BASE, "DevOps_Project_Report_NimmaYatri.pdf")

def pic(name):
    return os.path.join(PICS, name)

# ─────────────────────────────────────────────
# COLOURS
# ─────────────────────────────────────────────
C_BLACK    = RGBColor(0,   0,   0)
C_NAVY     = RGBColor(31,  78, 121)
C_WHITE    = RGBColor(255, 255, 255)
C_CAPTION  = RGBColor(80,  80,  80)
C_IMGBDR   = "AAAAAA"   # hex for image border
C_TBLHDR   = "1F4E79"   # hex for table header fill
C_TBLALT   = "EBF3FB"   # hex for alternate row fill
C_CODEBG   = "F5F5F5"   # hex for code block background

# Page dimensions (inches)
PAGE_W      = 8.5
MARGIN_L    = 1.25
MARGIN_R    = 1.0
MARGIN_T    = 1.0
MARGIN_B    = 1.0
TEXT_W      = PAGE_W - MARGIN_L - MARGIN_R   # 6.25"
IMG_W       = 5.5    # image width (fits safely inside text width)
TABLE_W     = 6.0    # max total table width

# ─────────────────────────────────────────────
# DOCUMENT SETUP
# ─────────────────────────────────────────────
doc = Document()

for sec in doc.sections:
    sec.top_margin    = Inches(MARGIN_T)
    sec.bottom_margin = Inches(MARGIN_B)
    sec.left_margin   = Inches(MARGIN_L)
    sec.right_margin  = Inches(MARGIN_R)
    sec.page_width    = Inches(PAGE_W)
    sec.page_height   = Inches(11)

# Normal (body) style
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing      = 1.5
pf.space_before      = Pt(0)
pf.space_after       = Pt(6)

# ─── Heading styles ──────────────────────────
HEADING_CONFIG = {
    1: dict(size=16, bold=True, caps=True,  align=WD_ALIGN_PARAGRAPH.CENTER, sb=0,  sa=10),
    2: dict(size=14, bold=True, caps=False, align=WD_ALIGN_PARAGRAPH.LEFT,   sb=12, sa=4),
    3: dict(size=12, bold=True, caps=False, align=WD_ALIGN_PARAGRAPH.LEFT,   sb=8,  sa=3),
}

def _set_style_font_xml(style_element, name):
    rPr = style_element.get_or_add_rPr()
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'),   name)
    rf.set(qn('w:hAnsi'),   name)
    rf.set(qn('w:eastAsia'),name)
    rPr.insert(0, rf)

for lvl, cfg in HEADING_CONFIG.items():
    st = doc.styles[f'Heading {lvl}']
    st.font.name      = 'Times New Roman'
    st.font.size      = Pt(cfg['size'])
    st.font.bold      = cfg['bold']
    st.font.italic    = False
    st.font.all_caps  = cfg['caps']
    st.font.color.rgb = C_BLACK
    st.paragraph_format.alignment          = cfg['align']
    st.paragraph_format.space_before       = Pt(cfg['sb'])
    st.paragraph_format.space_after        = Pt(cfg['sa'])
    st.paragraph_format.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
    st.paragraph_format.line_spacing       = 1.5
    st.paragraph_format.keep_with_next     = True
    _set_style_font_xml(st.element, 'Times New Roman')

# ─────────────────────────────────────────────
# HEADER  (project title on every page except title page)
# ─────────────────────────────────────────────
def setup_header_footer(section, first_page_special=False):
    """Add running header and page-number footer."""
    if first_page_special:
        spr = section._sectPr
        tpg = OxmlElement('w:titlePg')
        spr.append(tpg)

    # Header
    hdr = section.header
    hdr.is_linked_to_previous = False
    for p in hdr.paragraphs:
        p.clear()
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("Nimma Yatri – DevOps CI/CD Pipeline Report | Ramaiah Institute of Technology")
    hr.font.name  = 'Times New Roman'
    hr.font.size  = Pt(9)
    hr.font.color.rgb = RGBColor(120, 120, 120)

    # Horizontal rule under header
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Footer – page number
    ftr = section.footer
    ftr.is_linked_to_previous = False
    for p in ftr.paragraphs:
        p.clear()
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Rule above footer
    fpPr = fp._p.get_or_add_pPr()
    fpBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'AAAAAA')
    fpBdr.append(top)
    fpPr.append(fpBdr)

    def _field(run, code):
        fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
        it  = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = code
        fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate')
        fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end')
        for el in [fc1, it, fc2, fc3]:
            run._r.append(el)

    fr = fp.add_run("Page ")
    fr.font.name = 'Times New Roman'; fr.font.size = Pt(10)
    fr2 = fp.add_run()
    fr2.font.name = 'Times New Roman'; fr2.font.size = Pt(10)
    _field(fr2, ' PAGE ')
    fr3 = fp.add_run(" of ")
    fr3.font.name = 'Times New Roman'; fr3.font.size = Pt(10)
    fr4 = fp.add_run()
    fr4.font.name = 'Times New Roman'; fr4.font.size = Pt(10)
    _field(fr4, ' NUMPAGES ')

setup_header_footer(doc.sections[0], first_page_special=True)

# ─────────────────────────────────────────────
# LOW-LEVEL HELPERS
# ─────────────────────────────────────────────
def _fix_run_font(run, size=12, bold=False, italic=False, color=C_BLACK, all_caps=False):
    run.font.name      = 'Times New Roman'
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.all_caps  = all_caps
    run.font.color.rgb = color
    r  = run._r
    rPr = r.get_or_add_rPr()
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'),  'Times New Roman')
    rf.set(qn('w:hAnsi'),  'Times New Roman')
    rPr.insert(0, rf)

def _set_para_fmt(para, before=0, after=6, line=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = para.paragraph_format
    pf.space_before      = Pt(before)
    pf.space_after       = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = line
    para.alignment       = align

def _hex_fill(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ─────────────────────────────────────────────
# HIGH-LEVEL HELPERS
# ─────────────────────────────────────────────
def add_page_break(doc):
    doc.add_page_break()

def add_heading1(doc, text):
    p = doc.add_heading(text, level=1)
    _set_para_fmt(p, before=0, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        _fix_run_font(r, size=16, bold=True, all_caps=True)
    return p

def add_heading2(doc, text):
    p = doc.add_heading(text, level=2)
    _set_para_fmt(p, before=12, after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
    for r in p.runs:
        _fix_run_font(r, size=14, bold=True)
    return p

def add_heading3(doc, text):
    p = doc.add_heading(text, level=3)
    _set_para_fmt(p, before=8, after=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    for r in p.runs:
        _fix_run_font(r, size=12, bold=True)
    return p

def add_body(doc, text, justify=True):
    p = doc.add_paragraph()
    _set_para_fmt(p, before=0, after=6,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(text)
    _fix_run_font(r, size=12)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    _set_para_fmt(p, before=0, after=3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    if p.runs:
        p.runs[0].text = text
        _fix_run_font(p.runs[0], size=12)
    else:
        r = p.add_run(text)
        _fix_run_font(r, size=12)
    return p

def add_spacer(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)

def add_code_block(doc, code):
    p = doc.add_paragraph()
    _set_para_fmt(p, before=4, after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
    # Gray background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  C_CODEBG)
    pPr.append(shd)
    # Thin border
    pBdr = OxmlElement('w:pBdr')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '1')
        b.set(qn('w:color'), 'CCCCCC')
        pBdr.append(b)
    pPr.append(pBdr)
    r = p.add_run(code)
    r.font.name  = 'Courier New'
    r.font.size  = Pt(8.5)
    r.font.bold  = False
    r.font.color.rgb = C_BLACK
    return p

def add_image(doc, path, caption=""):
    """Image wrapped in a bordered single-cell table, centered, with italic caption."""
    add_spacer(doc, 4)
    if not os.path.exists(path):
        add_body(doc, f"[Image not found: {os.path.basename(path)}]")
        return

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = tbl.rows[0].cells[0]
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Cell inner padding
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top','72'),('left','72'),('bottom','72'),('right','72')]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    val)
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

    # Border
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '6')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), C_IMGBDR)
        tcBorders.append(b)
    tcPr.append(tcBorders)

    # Very light fill
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'FAFAFA')
    tcPr.append(shd)

    # Image
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    try:
        para.add_run().add_picture(path, width=Inches(IMG_W))
    except Exception as e:
        para.add_run(f"[Error: {e}]")

    # Caption
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(3)
        cp.paragraph_format.space_after  = Pt(10)
        cr = cp.add_run(caption)
        cr.font.name      = 'Times New Roman'
        cr.font.size      = Pt(10)
        cr.font.italic    = True
        cr.font.color.rgb = C_CAPTION
    else:
        add_spacer(doc, 8)

def add_table(doc, headers, rows, col_widths):
    """Styled table. col_widths must sum to ≤ TABLE_W (6.0")."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style     = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hcells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hcells[i].text = ""
        hp = hcells[i].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_fmt(hp, before=2, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        hr = hp.add_run(h)
        _fix_run_font(hr, size=12, bold=True, color=C_WHITE)
        _hex_fill(hcells[i], C_TBLHDR)
        if i < len(col_widths):
            hcells[i].width = Inches(col_widths[i])

    # Data rows
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            cp = cells[ci].paragraphs[0]
            _set_para_fmt(cp, before=2, after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
            cr = cp.add_run(val)
            _fix_run_font(cr, size=12)
            if ri % 2 == 0:
                _hex_fill(cells[ci], C_TBLALT)
            if ci < len(col_widths):
                cells[ci].width = Inches(col_widths[ci])
    add_spacer(doc, 6)
    return tbl

# ─────────────────────────────────────────────
# TITLE-PAGE HELPER
# ─────────────────────────────────────────────
def cp(doc, text, size=12, bold=False, before=4, after=4,
       color=C_BLACK, align=WD_ALIGN_PARAGRAPH.CENTER, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing  = 1.5
    r = p.add_run(text)
    _fix_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p

# ═══════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════
cp(doc, "RAMAIAH INSTITUTE OF TECHNOLOGY", 18, True, before=50, after=4, color=C_NAVY)
cp(doc, "Department of Information Science and Engineering", 13, True, before=2, after=2)
cp(doc, "Bengaluru – 560 054", 12, False, before=2, after=20)

# Divider line
div = doc.add_paragraph()
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
div.paragraph_format.space_before = Pt(0)
div.paragraph_format.space_after  = Pt(0)
pPr = div._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
for side in ['top','bottom']:
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'),   'single')
    b.set(qn('w:sz'),    '8')
    b.set(qn('w:space'), '4')
    b.set(qn('w:color'), '1F4E79')
    pBdr.append(b)
pPr.append(pBdr)

cp(doc, "DevOps Laboratory Project Report", 20, True, before=18, after=6, color=C_NAVY)
cp(doc, "Nimma Yatri – Bengaluru Auto-Rickshaw Survival Tool", 15, True, before=2, after=4, color=C_NAVY)
cp(doc, "A Complete CI/CD Pipeline Implementation using Jenkins, ESLint, Trivy, Docker and Vercel",
   12, False, before=0, after=24)

cp(doc, "Submitted by", 12, False, before=10, after=6)

# Student table
st = doc.add_table(rows=4, cols=2)
st.alignment = WD_TABLE_ALIGNMENT.CENTER
details = [
    ("Student Name",   "Harshendra M"),
    ("USN",            "1MS22IS053"),
    ("Section",        "IS-A  |  4th Semester"),
    ("Subject",        "DevOps Laboratory (21ISL67)"),
]
for i,(lbl,val) in enumerate(details):
    lc = st.rows[i].cells[0]; vc = st.rows[i].cells[1]
    lp = lc.paragraphs[0]; vp = vc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run(lbl + " :"); _fix_run_font(lr, 12, bold=True)
    vr = vp.add_run(val);        _fix_run_font(vr, 12)
    lc.width = Inches(2.2); vc.width = Inches(3.5)

cp(doc, "", 12, before=20, after=4)
cp(doc, "Academic Year: 2025 – 2026", 12, False, before=4, after=2)
cp(doc, "Date of Submission: May 2026", 12, False, before=2, after=4)

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# TABLE OF CONTENTS  (Word auto-field, updated by win32com)
# ═══════════════════════════════════════════════════════
toc_title = doc.add_paragraph()
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_title.paragraph_format.space_before = Pt(0)
toc_title.paragraph_format.space_after  = Pt(16)
tr = toc_title.add_run("TABLE OF CONTENTS")
_fix_run_font(tr, size=16, bold=True)

# Insert TOC field
toc_para = doc.add_paragraph()
toc_para.paragraph_format.space_before = Pt(0)
toc_para.paragraph_format.space_after  = Pt(0)
run = toc_para.add_run()

for fc_type in ['begin', 'separate', 'end']:
    if fc_type == 'begin':
        fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin')
        run._r.append(fc)
        it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
        it.text = r' TOC \o "1-3" \h \z \u '
        run._r.append(it)
    elif fc_type == 'separate':
        fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'separate')
        run._r.append(fc)
        ph = OxmlElement('w:t')
        ph.text = '[ Right-click here in Word and choose "Update Field" to populate the Table of Contents ]'
        run._r.append(ph)
    else:
        fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'end')
        run._r.append(fc)

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# CH 1: ABSTRACT
# ═══════════════════════════════════════════════════════
add_heading1(doc, "Chapter 1 – Abstract")
add_spacer(doc, 4)

add_body(doc,
    "This report presents the design, implementation, and evaluation of a complete DevOps CI/CD "
    "pipeline built around Nimma Yatri — a production-grade, AI-powered Next.js web application "
    "that helps Bengaluru commuters navigate auto-rickshaw fare disputes. The primary purpose of "
    "this project is to demonstrate end-to-end automation of software delivery using "
    "industry-standard DevOps tools, integrating source code management, automated testing, "
    "static code analysis, vulnerability scanning, containerisation, and cloud deployment within "
    "a single, cohesive pipeline."
)
add_body(doc,
    "The technology stack employed includes GitHub for source code management and webhook-based "
    "change detection, Jenkins 2.541.2 as the CI/CD automation server, ESLint for static code "
    "quality analysis, Trivy 0.69.3 for both filesystem and Docker image vulnerability scanning, "
    "Docker Desktop 28.0.4 for containerising the Next.js application, and Docker Hub as the "
    "container image registry. The deployment target is Vercel, which provides seamless cloud "
    "hosting for the Next.js application at the public URL https://nimmayatri.vercel.app."
)
add_body(doc,
    "The CI/CD workflow is triggered automatically when a developer pushes code to the main "
    "branch on GitHub. A GitHub webhook, tunnelled through ngrok, notifies Jenkins instantly, "
    "triggering the pipeline. Stages execute sequentially: source checkout, dependency "
    "installation, ESLint code quality check, Next.js production build, Trivy filesystem scan, "
    "Docker image build, Trivy image scan, Docker Hub login, and Docker push. A secondary "
    "no-push pipeline variant was also configured for development testing."
)
add_body(doc,
    "The outcome confirmed successful automation: the Docker image "
    "suicide768/nimmayatri-app:latest was published to Docker Hub, both Trivy reports were "
    "archived as Jenkins build artefacts (trivy-fs-report.txt: 31.66 KB; trivy-image-report.txt: "
    "212.44 KB), and the pipeline completed in under 4 minutes, with zero ESLint errors "
    "reported. The full pipeline was triggered automatically via GitHub webhook push by "
    "harshendram, confirming end-to-end DevOps automation."
)

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# CH 2: OBJECTIVE
# ═══════════════════════════════════════════════════════
add_heading1(doc, "Chapter 2 – Objective")
add_spacer(doc, 4)

add_heading2(doc, "2.1  Why CI/CD Pipelines are Important")
add_body(doc,
    "Continuous Integration and Continuous Deployment (CI/CD) pipelines are the backbone of "
    "modern software engineering. Traditional software delivery — where integration, testing, "
    "and deployment are manual, periodic activities — leads to 'integration hell', long release "
    "cycles, and unreliable deployments. CI/CD pipelines automate the entire lifecycle from "
    "code commit to production, ensuring every change is built, validated, and deployed "
    "consistently and quickly. This eliminates manual errors, reduces risk, and provides "
    "immediate feedback to developers within minutes of a commit."
)

add_heading2(doc, "2.2  Goals of Automation")
add_body(doc, "The specific automation goals of this project are:")
add_bullet(doc, "Eliminate all manual build and deployment steps through a Jenkins pipeline triggered automatically by a GitHub push event.")
add_bullet(doc, "Enforce code quality standards at every commit using ESLint, blocking poorly formatted or non-compliant code from advancing in the pipeline.")
add_bullet(doc, "Detect security vulnerabilities early by scanning both the npm dependency tree and the Docker container image with Trivy before code reaches production.")
add_bullet(doc, "Standardise deployment environments using Docker multi-stage builds, guaranteeing identical behaviour across development, CI, and production.")
add_bullet(doc, "Automate container image versioning (using Jenkins build numbers) and publishing to Docker Hub for every successful build.")
add_bullet(doc, "Enable zero-touch execution using GitHub Webhooks, so a git push is the only action required from the developer.")

add_heading2(doc, "2.3  Benefits of DevOps Practices")
add_body(doc, "Adopting a DevOps culture delivers the following measurable benefits:")
add_bullet(doc, "Faster Time to Market: Automated pipelines reduce deployment cycles from days to minutes, enabling frequent and reliable releases.")
add_bullet(doc, "Improved Code Quality: Automated linting and build checks catch errors before they reach production, reducing post-release defect rates.")
add_bullet(doc, "Enhanced Security Posture: Continuous vulnerability scanning with Trivy identifies known CVEs in dependencies and container images, enabling timely patching.")
add_bullet(doc, "Environment Consistency: Docker containerisation eliminates the 'works on my machine' problem, ensuring the application runs identically everywhere.")
add_bullet(doc, "Full Traceability: Every pipeline execution is linked to a specific git commit, providing complete audit trail from code change to deployed artefact.")
add_bullet(doc, "Reduced Deployment Risk: Consistent, automated build processes and Docker image versioning enable rapid rollback to any previous build.")

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# CH 3: TOOLS AND TECHNOLOGIES
# ═══════════════════════════════════════════════════════
add_heading1(doc, "Chapter 3 – Tools and Technologies Used")
add_spacer(doc, 4)
add_body(doc, "The following table summarises all tools and technologies used in this project:")
add_spacer(doc, 4)

add_table(doc,
    headers=["Tool / Technology", "Purpose in the Pipeline"],
    rows=[
        ["GitHub",                   "Source code management, version control, and webhook-based change notification to Jenkins."],
        ["Jenkins 2.541.2",          "CI/CD automation server orchestrating all pipeline stages from checkout to Docker push."],
        ["ESLint (eslint-config-next)","Static code quality analysis enforcing Next.js core-web-vitals rules across the TypeScript codebase."],
        ["Trivy 0.69.3",             "Open-source vulnerability scanner for filesystem dependencies (npm) and Docker container images."],
        ["Docker Desktop 28.0.4",    "Containerisation platform used to build a production-ready multi-stage Docker image."],
        ["Docker Hub",               "Cloud container image registry storing versioned suicide768/nimmayatri-app images."],
        ["Node.js 18 (LTS)",         "JavaScript runtime for npm install, ESLint checks, and Next.js production builds."],
        ["Next.js 14.2.21",          "React-based full-stack web framework forming the application codebase."],
        ["ngrok",                    "Secure tunnel exposing the local Jenkins server to GitHub for webhook delivery."],
        ["Vercel",                   "Cloud deployment platform hosting the live Next.js application at nimmayatri.vercel.app."],
        ["Git 2.44.0",               "Distributed version control system for all source code tracking."],
        ["npm",                      "Node.js package manager for dependency installation and build script execution."],
    ],
    col_widths=[2.0, 4.0]
)

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# CH 4: PROJECT TECH STACK
# ═══════════════════════════════════════════════════════
add_heading1(doc, "Chapter 4 – Project Tech Stack")
add_spacer(doc, 4)

add_heading2(doc, "4.1  Project Overview")
add_body(doc,
    "Nimma Yatri (meaning 'Your Journey' in Kannada) is an AI-powered web application designed "
    "to help Bengaluru commuters navigate the notorious auto-rickshaw fare dispute culture. "
    "The application provides AI-assisted fare calculation, real-time voice/video negotiation "
    "support, Kannada phrase flashcards, a Scam-O-Meter gauge, a Panic Button safety feature, "
    "and community intelligence from r/bangalore. It was built for the Kiro Week 5 Challenge "
    "and deployed live at https://nimmayatri.vercel.app."
)
add_body(doc,
    "The application serves as the real-world, production-grade subject for this DevOps pipeline "
    "implementation. With 12 React components, 4 API routes, 4 custom hooks, and 6 utility "
    "libraries comprising over 8,500 lines of TypeScript, it represents a non-trivial codebase "
    "on which the full DevOps toolchain is exercised."
)

add_heading2(doc, "4.2  Frontend Architecture")
add_body(doc,
    "The frontend is built on Next.js 14 App Router with TypeScript. UI components include: "
    "FareCalculator (Google Maps integration + Scam-O-Meter speedometer gauge), "
    "FloatingLiveAssistant (real-time voice/video via Gemini 2.5 Flash WebSocket), "
    "KannadaPhrases (audio phrase deck with Text-to-Speech), PanicButton (emergency safety "
    "feature), Chatbot (AI text assistant), and RedditPosts (r/bangalore community feed). "
    "Styling uses Tailwind CSS 3.4 with Framer Motion animations. Multi-language support for "
    "10 Indian languages is handled through a React Context provider."
)

add_heading2(doc, "4.3  Backend and API Routes")
add_body(doc,
    "Next.js API routes provide the backend layer. The /api/chatbot endpoint integrates Google "
    "Gemini 2.5 Flash with model fallback (Gemini 2.0). The /api/fare endpoint computes fares "
    "using government meter rates (₹30 base + ₹15/km, 1.5× night rate). The /api/gemini-live "
    "endpoint manages WebSocket configuration for multimodal AI sessions. Google Maps Platform "
    "APIs handle all location services."
)

add_heading2(doc, "4.4  Technology Stack Summary")
add_spacer(doc, 4)
add_table(doc,
    headers=["Layer", "Technology", "Version"],
    rows=[
        ["Framework",       "Next.js (App Router)",                    "14.2.21"],
        ["Language",        "TypeScript",                              "5.7.2"],
        ["Styling",         "Tailwind CSS + Framer Motion",            "3.4.17 / 11.15.0"],
        ["AI Engine",       "Google Gemini Multimodal Live API",       "2.5 Flash"],
        ["Maps",            "Google Maps Platform",                    "Places + Distance Matrix"],
        ["State",           "React Context API",                       "18.3.1"],
        ["Icons",           "Lucide React",                            "0.468.0"],
        ["Runtime",         "Node.js",                                 "18 LTS"],
        ["Container",       "Docker (multi-stage build)",              "28.0.4"],
        ["Deployment",      "Vercel – Mumbai (bom1)",                  "Latest"],
        ["Performance",     "Lighthouse Score",                        "94 / 100"],
    ],
    col_widths=[1.7, 2.8, 1.5]
)

add_heading2(doc, "4.5  Application Screenshots")
add_image(doc, pic("Screenshot 2026-05-18 184426.png"),
    "Fig 4.1 – Nimma Yatri running on localhost:3000 showing the Fare Calculator (left) "
    "and Kannada Phrases flashcard deck (right)")
add_image(doc, pic("Screenshot 2026-05-18 183526.png"),
    "Fig 4.2 – VS Code terminal: npm run lint returns '✓ No ESLint warnings or errors'; "
    "npm run build completes the Next.js production build successfully; Docker version confirmed")

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# CH 5: GITHUB REPOSITORY DETAILS
# ═══════════════════════════════════════════════════════
add_heading1(doc, "Chapter 5 – GitHub Repository Details")
add_spacer(doc, 4)

add_heading2(doc, "5.1  Repository Information")
add_spacer(doc, 2)
add_table(doc,
    headers=["Attribute", "Value"],
    rows=[
        ["Repository Name",  "Devops_Lab_Assign"],
        ["Owner",            "harshendram (Harshendra M)"],
        ["Clone URL",        "https://github.com/harshendram/Devops_Lab_Assign.git"],
        ["Primary Branch",   "main"],
        ["Visibility",       "Public"],
        ["Language Composition", "97.1% TypeScript, 2.5% CSS"],
        ["Total Deployments","6 (Vercel)"],
        ["Contributors",     "1 – harshendram"],
    ],
    col_widths=[2.0, 4.0]
)

add_heading2(doc, "5.2  Repository Structure")
add_body(doc, "Key files and directories relevant to the DevOps pipeline:")
add_code_block(doc,
"""nimmayatri/
├── app/                      # Next.js App Router pages + API routes
│   └── api/                  # chatbot / fare / gemini / gemini-live
├── components/               # 12 React UI components
├── context/ hooks/ lib/      # State, custom hooks, utilities
├── public/assets/            # Static assets
├── Jenkinsfile               # Full CI/CD pipeline (with Docker push)
├── Jenkinsfile-no-push       # Pipeline variant without Docker push
├── Dockerfile                # Multi-stage Docker build (3 stages)
├── .eslintrc.json            # ESLint config (next/core-web-vitals)
├── next.config.mjs           # Next.js config (standalone output)
├── package.json              # Dependencies and build scripts
└── README.md                 # Project documentation""")

add_heading2(doc, "5.3  Branching Strategy")
add_body(doc,
    "The project follows a single-branch strategy with main as the only active branch. "
    "Since this is a single-developer lab project, multi-branch GitFlow is not required. "
    "All development commits and pipeline triggers operate on main. Every push to main "
    "immediately activates the Jenkins pipeline via the configured GitHub webhook."
)

add_heading2(doc, "5.4  GitHub Repository Screenshot")
add_image(doc, pic("Screenshot 2026-05-22 120841.png"),
    "Fig 5.1 – GitHub repository root view: commit 'testing for vercel deployment' by harshendram "
    "(cb350df), complete project file tree, 6 Vercel deployments, single contributor, "
    "97.1% TypeScript language composition")

add_heading2(doc, "5.5  GitHub Webhook Configuration")
add_body(doc,
    "A GitHub webhook was configured to notify Jenkins whenever code is pushed to the repository. "
    "Because Jenkins runs on a local Windows machine without a public IP, ngrok was used to "
    "create a secure HTTPS tunnel, exposing the Jenkins /github-webhook/ endpoint to the internet."
)
add_heading3(doc, "5.5.1  Setup Steps")
add_bullet(doc, "Step 1: Start ngrok: ngrok http 8080 – generates a public HTTPS URL (e.g., https://semiobliviously-unevaporative-neriah.ngrok-free.dev).")
add_bullet(doc, "Step 2: In GitHub → Settings → Webhooks → Add webhook.")
add_bullet(doc, "Step 3: Set Payload URL to: https://<ngrok-url>/github-webhook/")
add_bullet(doc, "Step 4: Content type: application/json; Trigger: Just the push event.")
add_bullet(doc, "Step 5: Click 'Add webhook'. A green tick confirms successful delivery of the test ping.")
add_bullet(doc, "Step 6: In Jenkins, enable 'GitHub hook trigger for GITScm polling' in the job's Build Triggers section.")

add_image(doc, pic("Screenshot 2026-05-18 210619.png"),
    "Fig 5.2 – GitHub Webhooks / Manage Webhook page: ngrok-generated Payload URL configured, "
    "Content type = application/json, SSL verification enabled, trigger = push event only")
add_image(doc, pic("Screenshot 2026-05-18 212826.png"),
    "Fig 5.3 – ngrok session dashboard: Forwarding URL → http://127.0.0.1:8080, Region: India, "
    "HTTP request log showing POST /github-webhook/ → 200 OK, confirming successful webhook delivery to Jenkins")

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# CH 6: JENKINS CONFIGURATION
# ═══════════════════════════════════════════════════════
add_heading1(doc, "Chapter 6 – Jenkins Configuration")
add_spacer(doc, 4)

add_heading2(doc, "6.1  Jenkins Installation")
add_body(doc,
    "Jenkins 2.541.2 was installed on a Windows 11 machine using the official Jenkins LTS "
    "Windows installer. The installer registers Jenkins as a Windows Service accessible at "
    "http://localhost:8080. Initial setup involved unlocking Jenkins with the auto-generated "
    "administrator password, installing suggested plugins, and creating an admin user account."
)

add_heading2(doc, "6.2  Required Plugins")
add_spacer(doc, 2)
add_table(doc,
    headers=["Plugin", "Purpose"],
    rows=[
        ["Pipeline",                "Enables declarative Jenkinsfile pipeline support."],
        ["Git Plugin",              "Clones and polls GitHub repositories."],
        ["GitHub Integration Plugin","Enables webhook-triggered builds from GitHub push events."],
        ["NodeJS Plugin",           "Provides Node.js 18 tool (configured as 'Node18') for npm commands."],
        ["Credentials Binding",     "Securely injects Docker Hub credentials into pipeline stages."],
        ["Pipeline: Stage View",    "Visual stage-by-stage pipeline execution dashboard."],
        ["Workspace Cleanup Plugin","Cleans workspace between builds for consistency."],
    ],
    col_widths=[2.2, 3.8]
)

add_heading2(doc, "6.3  Pipeline Jobs Created")
add_body(doc, "Two pipeline jobs were created in Jenkins:")
add_bullet(doc, "nimmayatri-pipeline: Full CI/CD pipeline including Docker Hub login and image push. Uses Jenkinsfile from the repository root.")
add_bullet(doc, "nimmayatri-pipeline-no-push: Testing variant executing all stages except Docker Hub login and push. Uses Jenkinsfile-no-push.")
add_body(doc, "Each job is configured as a 'Pipeline' type with: Definition = Pipeline script from SCM; SCM = Git (main branch); GitHub hook trigger enabled.")

add_heading2(doc, "6.4  Jenkinsfile Explanation – Full Pipeline")
add_body(doc, "The primary Jenkinsfile defines a declarative pipeline with 9 stages:")
add_code_block(doc,
"""pipeline {
    agent any
    tools { nodejs 'Node18' }
    environment {
        IMAGE_NAME = "suicide768/nimmayatri-app"
        IMAGE_TAG  = "${env.BUILD_NUMBER}"
    }
    stages {
        stage('Checkout')               { steps { git branch: 'main', url: '...' } }
        stage('Install Dependencies')   { steps { bat 'npm install' } }
        stage('ESLint Check')           { steps { bat 'npm run lint' } }
        stage('Build Next.js App')      { steps { bat 'npm run build' } }
        stage('Trivy Filesystem Scan')  { steps { bat 'trivy fs . > trivy-fs-report.txt' } }
        stage('Docker Build')           { steps { bat "docker build -t ..." } }
        stage('Trivy Docker Image Scan'){ steps { bat "trivy image ... > trivy-image-report.txt" } }
        stage('Docker Hub Login')       { steps { withCredentials([...]) {...} } }
        stage('Docker Push')            { steps { bat "docker push ..." } }
    }
    post {
        always  { archiveArtifacts artifacts: 'trivy-*.txt', allowEmptyArchive: true }
        success { echo 'Pipeline executed successfully!' }
        failure { echo 'Pipeline failed!' }
    }
}""")

add_heading3(doc, "6.4.1  Stage-by-Stage Explanation")
add_body(doc, "agent any: Runs the pipeline on any available Jenkins build node — in this case, the local Windows controller.")
add_body(doc, "tools { nodejs 'Node18' }: Ensures Node.js 18 is on the PATH before any stage executes, enabling all npm and node commands.")
add_body(doc, "environment block: Defines IMAGE_NAME = 'suicide768/nimmayatri-app' and IMAGE_TAG = BUILD_NUMBER (Jenkins build number). Using the build number as a tag provides automatic immutable versioning: build 1 → :1, build 2 → :2, etc., enabling rollback to any previous build.")
add_body(doc, "Checkout: Clones the main branch from GitHub. When webhook-triggered, Jenkins fetches the exact commit that triggered the build.")
add_body(doc, "Install Dependencies: Runs npm install, installing all packages from package.json including runtime and dev dependencies into node_modules.")
add_body(doc, "ESLint Check: Executes npm run lint (next lint). If any ESLint error is found, the stage fails immediately, halting the pipeline and preventing non-compliant code from advancing.")
add_body(doc, "Build Next.js App: Runs npm run build (next build), compiling TypeScript, performing type checking, generating optimised JS bundles, and creating the .next/standalone output directory — a self-contained Next.js server required for Docker deployment.")
add_body(doc, "Trivy Filesystem Scan: Runs trivy fs . to scan package-lock.json and node_modules for CVEs using the NVD and GitHub Advisory databases. Output saved to trivy-fs-report.txt.")
add_body(doc, "Docker Build: Builds the multi-stage Docker image with two tags: the build-number tag (e.g., :2) and :latest.")
add_body(doc, "Trivy Docker Image Scan: Scans the built Docker image for OS-level (Alpine) and application-level CVEs. Output saved to trivy-image-report.txt.")
add_body(doc, "Docker Hub Login: Uses withCredentials to inject Docker Hub credentials stored securely in Jenkins (ID: dockerhub-creds) as environment variables, then runs docker login via stdin to avoid credentials appearing in logs.")
add_body(doc, "Docker Push: Pushes both the build-number tag and :latest to Docker Hub, making the containerised application publicly accessible.")
add_body(doc, "post/always block: Archives trivy-*.txt files as Jenkins build artefacts regardless of pipeline success or failure, providing permanent security scan records.")

add_heading2(doc, "6.5  Jenkinsfile-no-push Explanation")
add_body(doc,
    "The Jenkinsfile-no-push is identical to the full pipeline except the Docker Hub Login and "
    "Docker Push stages are removed. This variant was used during development to verify that "
    "build, lint, scan, and Docker build stages work correctly without requiring Docker Hub "
    "credentials. It is suitable for pull request validation or environments without registry access. "
    "Both variants archive Trivy reports as build artefacts."
)

add_heading2(doc, "6.6  Jenkins Pipeline Screenshots")
add_image(doc, pic("Screenshot 2026-05-18 212111.png"),
    "Fig 6.1 – Jenkins Dashboard: nimmayatri-pipeline (Last Success: 5 min 57 sec, Build #2) "
    "and nimmayatri-pipeline-no-push (Last Success: 59 min, Build #2), both showing green success status")
add_image(doc, pic("Screenshot 2026-05-22 120051.png"),
    "Fig 6.2 – Jenkins Builds history (user: Harshendra): all builds stable — "
    "nimmayatri-pipeline-no-push #1 and #2, nimmayatri-pipeline #1, firstjob #1 — all green checkmarks")
add_image(doc, pic("Screenshot 2026-05-22 120108.png"),
    "Fig 6.3 – Jenkins project page for nimmayatri-pipeline-no-push: Last Successful Artefacts — "
    "trivy-fs-report.txt (31.66 KiB) and trivy-image-report.txt (212.44 KiB); "
    "Build #2 = Last stable, Last successful, Last completed build (May 18, 2026 at 8:21 PM)")
add_image(doc, pic("Screenshot 2026-05-18 212056.png"),
    "Fig 6.4 – Jenkins Build #2 status page: 'Started by GitHub push by harshendram' — "
    "confirms automatic webhook-triggered execution; build took 3 min 54 sec")
add_image(doc, pic("Screenshot 2026-05-18 202304.png"),
    "Fig 6.5 – Jenkins Build #2 details: Build Artefacts section showing trivy-fs-report.txt (31.66 KB) "
    "and trivy-image-report.txt (212.64 KB), confirming both Trivy scans completed and archived")
add_image(doc, pic("Screenshot 2026-05-18 212021.png"),
    "Fig 6.6 – Jenkins Console Output (top): 'Started by GitHub push by harshendram', "
    "Jenkinsfile retrieved from git, pipeline stages beginning execution")
add_image(doc, pic("Screenshot 2026-05-18 212040.png"),
    "Fig 6.7 – Jenkins Console Output (tail): Docker layers pushed, artefacts archived, "
    "'Pipeline executed successfully!' — Finished: SUCCESS")

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# CH 7: CODE QUALITY ANALYSIS
# ═══════════════════════════════════════════════════════
add_heading1(doc, "Chapter 7 – Code Quality Analysis")
add_spacer(doc, 4)

add_heading2(doc, "7.1  Tool Used – ESLint")
add_body(doc,
    "ESLint is the industry-standard JavaScript and TypeScript static analysis tool. For this "
    "project, ESLint runs via Next.js's built-in lint integration (next lint), which wraps ESLint "
    "with Next.js-specific rule sets. The configuration file .eslintrc.json extends "
    "next/core-web-vitals — the strictest built-in Next.js preset, covering React best practices, "
    "Hooks rules, accessibility, and Core Web Vitals performance patterns."
)

add_heading2(doc, "7.2  Configuration")
add_code_block(doc,
"""{
  "extends": "next/core-web-vitals",
  "rules": {
    "react/no-unescaped-entities": "off",
    "@next/next/no-img-element": "off"
  }
}""")
add_body(doc, "The next/core-web-vitals preset includes:")
add_bullet(doc, "eslint:recommended – Standard JS best practices (no-unused-vars, no-undef, etc.).")
add_bullet(doc, "plugin:react/recommended – React-specific rules (prop-types, JSX correctness).")
add_bullet(doc, "plugin:react-hooks/recommended – Rules of Hooks (exhaustive-deps, rules-of-hooks).")
add_bullet(doc, "@next/next/core-web-vitals – Performance rules for Next.js-specific patterns.")
add_bullet(doc, "plugin:jsx-a11y/recommended – Accessibility rules for JSX elements.")
add_body(doc, "Two rules were explicitly disabled:")
add_bullet(doc, "react/no-unescaped-entities – Disabled to accommodate Kannada text strings and apostrophes in JSX content.")
add_bullet(doc, "@next/next/no-img-element – Disabled for specific components where the Next.js Image optimisation component is not applicable.")

add_heading2(doc, "7.3  ESLint in the Pipeline")
add_body(doc,
    "ESLint runs as Stage 3 of the Jenkins pipeline (npm run lint) after dependency installation "
    "but before the production build. If any ESLint error is detected, the stage exits with a "
    "non-zero code, Jenkins marks the build as FAILED, and all subsequent stages are skipped. "
    "This ensures no non-compliant code advances to build, scan, or deployment stages."
)

add_heading2(doc, "7.4  Analysis Results")
add_body(doc,
    "The ESLint analysis of the complete Nimma Yatri codebase returned zero warnings and zero "
    "errors across all 12 React components, 4 API routes, 4 custom hooks, and 6 utility "
    "libraries. Key validations confirmed by the clean lint result:"
)
add_bullet(doc, "No unused variables or imports across the TypeScript codebase.")
add_bullet(doc, "Correct React Hooks usage — no conditional hook calls, no missing dependency arrays.")
add_bullet(doc, "No deprecated Next.js patterns or performance anti-patterns.")
add_bullet(doc, "No JSX accessibility violations.")

add_image(doc, pic("Screenshot 2026-05-18 183526.png"),
    "Fig 7.1 – VS Code terminal: npm run lint returns '✓ No ESLint warnings or errors'; "
    "npm run build confirms a clean Next.js production build; Docker version 28.0.4 verified")

add_heading2(doc, "7.5  Issues Fixed During Development")
add_bullet(doc, "Unescaped Kannada entities in JSX: react/no-unescaped-entities errors caused by Kannada text. Resolution: rule disabled in .eslintrc.json.")
add_bullet(doc, "Standard <img> elements: @next/next/no-img-element warnings on specific components. Resolution: rule disabled where Next.js Image was not applicable.")
add_bullet(doc, "Incomplete Hook dependency arrays: react-hooks/exhaustive-deps warnings on useEffect hooks. Resolution: dependency arrays updated to include all referenced variables.")

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# CH 8: DEPENDENCY AND VULNERABILITY SCANNING
# ═══════════════════════════════════════════════════════
add_heading1(doc, "Chapter 8 – Dependency and Vulnerability Scanning")
add_spacer(doc, 4)

add_heading2(doc, "8.1  Tool Used – Trivy")
add_body(doc,
    "Trivy (version 0.69.3) is an open-source comprehensive vulnerability scanner by Aqua "
    "Security. It is the industry's most widely adopted container and filesystem security "
    "scanner, detecting CVEs in OS packages, language-specific dependencies (npm, pip, etc.), "
    "IaC misconfigurations, and secret leaks. Trivy was chosen for this project due to its "
    "seamless CI pipeline integration, broad vulnerability database coverage (NVD, GitHub "
    "Advisory Database, OS vendor advisories), and ability to scan both project filesystems "
    "and Docker images."
)

add_heading2(doc, "8.2  Scan Process")
add_body(doc, "Two Trivy scans are performed sequentially in the pipeline:")
add_bullet(doc, "Stage 5 – Filesystem Scan (trivy fs .): Scans the project directory focusing on package-lock.json for npm dependency CVEs. Runs after the Next.js build, while node_modules are present. Output saved to trivy-fs-report.txt.")
add_bullet(doc, "Stage 7 – Docker Image Scan (trivy image <image:tag>): Scans the built Docker image for vulnerabilities in Alpine Linux OS packages and Node.js runtime packages within the container. Output saved to trivy-image-report.txt.")

add_heading2(doc, "8.3  Filesystem Scan Results")
add_body(doc, "Trivy filesystem scan of package-lock.json identified 28 vulnerabilities:")
add_spacer(doc, 2)
add_table(doc,
    headers=["Severity", "Count"],
    rows=[
        ["CRITICAL", "1"],
        ["HIGH",     "9"],
        ["MEDIUM",   "14"],
        ["LOW",      "4"],
        ["UNKNOWN",  "0"],
        ["TOTAL",    "28"],
    ],
    col_widths=[2.5, 1.5]
)
add_body(doc, "Key findings — all in the next (Next.js 14.2.21) package:")
add_bullet(doc, "CVE-2025-29927 [CRITICAL] – Authorization Bypass in Next.js Middleware. Fixed in v14.2.25. Highest priority remediation.")
add_bullet(doc, "CVE-2026-44573 [HIGH] – Middleware/Proxy bypass in Pages Router applications.")
add_bullet(doc, "CVE-2026-44578 [HIGH] – Server-side request forgery via WebSocket upgrades.")
add_bullet(doc, "GHSA-5j59-xgg2-r9c4 [HIGH] – Denial of Service via Server Components (incomplete fix).")
add_bullet(doc, "GHSA-8h8q-6873-q5fj [HIGH] – Denial of Service via Server Components.")
add_bullet(doc, "CVE-2025-55173 [MEDIUM] – Content Injection vulnerability in Image Optimisation.")

add_image(doc, pic("Screenshot 2026-05-18 185812.png"),
    "Fig 8.1 – Trivy filesystem scan summary: package-lock.json — Total 28 vulnerabilities "
    "(0 UNKNOWN, 4 LOW, 14 MEDIUM, 9 HIGH, 1 CRITICAL)")
add_image(doc, pic("Screenshot 2026-05-18 185826.png"),
    "Fig 8.2 – Trivy CVE detail table: CVE-2025-29927 (CRITICAL), CVE-2026-44573 (HIGH), "
    "CVE-2026-44578 (HIGH) with descriptions and affected Next.js versions")
add_image(doc, pic("Screenshot 2026-05-18 185844.png"),
    "Fig 8.3 – Trivy scan continued: additional HIGH and MEDIUM severity CVEs in the "
    "Next.js dependency with full CVE identifiers and advisory links")

add_heading2(doc, "8.4  Docker Image Scan Results")
add_body(doc,
    "The Trivy Docker image scan of the nimmayatri-app (based on node:18-alpine / Alpine 3.21.3) "
    "identified approximately 58 vulnerabilities, including CVEs in Alpine Linux OS packages "
    "(BusyBox) and Node.js runtime in addition to npm packages:"
)
add_spacer(doc, 2)
add_table(doc,
    headers=["Severity", "Count"],
    rows=[
        ["CRITICAL", "4"],
        ["HIGH",     "15"],
        ["MEDIUM",   "26"],
        ["LOW",      "5"],
        ["UNKNOWN",  "0"],
        ["TOTAL",    "~58"],
    ],
    col_widths=[2.5, 1.5]
)
add_body(doc, "Additional CVEs beyond the filesystem scan (Alpine Linux packages):")
add_bullet(doc, "CVE-2024-58251 [MEDIUM] – BusyBox: local users can launch processes with altered network namespaces.")
add_bullet(doc, "CVE-2025-46394 [LOW] – BusyBox: TAR archive filename handling vulnerability.")
add_bullet(doc, "CVE-2024-58251 [MEDIUM] – BusyBox/binsh: same network namespace vulnerability.")

add_image(doc, pic("Screenshot 2026-05-18 190253.png"),
    "Fig 8.4 – Trivy filesystem scan: detailed node_modules package vulnerability list")
add_image(doc, pic("Screenshot 2026-05-18 190306.png"),
    "Fig 8.5 – Trivy Docker image scan summary for nimmayatri-app (Alpine 3.21.3): "
    "Total ~58 vulnerabilities including 4 CRITICAL, BusyBox and npm package findings")
add_image(doc, pic("Screenshot 2026-05-18 190521.png"),
    "Fig 8.6 – Trivy image scan output: OS family 'alpine' detected, scanning process, "
    "CVE-2026-31082 and related entries with severity levels")

add_heading2(doc, "8.5  Mitigation Steps")
add_bullet(doc, "Upgrade Next.js to ≥ 14.2.25 to resolve CVE-2025-29927 (CRITICAL). This is the highest priority action.")
add_bullet(doc, "Update the Docker base image from node:18-alpine to the latest patched version to resolve BusyBox CVEs.")
add_bullet(doc, "Run npm audit fix after the Next.js upgrade to resolve transitive dependency vulnerabilities automatically.")
add_bullet(doc, "Trivy reports are archived as Jenkins build artefacts for audit trail and periodic review.")
add_bullet(doc, "No secrets were detected by Trivy's secret scanner, confirming proper use of .env.local (gitignored) and Vercel environment variable configuration.")

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# CH 9: DOCKER HUB INTEGRATION
# ═══════════════════════════════════════════════════════
add_heading1(doc, "Chapter 9 – Docker Hub Integration")
add_spacer(doc, 4)

add_heading2(doc, "9.1  Dockerfile Explanation")
add_body(doc,
    "The Dockerfile uses a three-stage multi-stage build to produce a minimal, secure, "
    "production-ready image. Multi-stage builds keep build tools out of the final image, "
    "significantly reducing image size and attack surface."
)
add_code_block(doc,
"""# ── Stage 1: deps ─────────────────────────────────────────────
FROM node:18-alpine AS deps
WORKDIR /app
COPY package.json package-lock.json* ./
RUN npm install

# ── Stage 2: builder ───────────────────────────────────────────
FROM node:18-alpine AS builder
WORKDIR /app
COPY --from=deps /app/node_modules ./node_modules
COPY . .
RUN npm run build

# ── Stage 3: runner (final production image) ──────────────────
FROM node:18-alpine AS runner
WORKDIR /app
ENV NODE_ENV=production

RUN addgroup -S nodejs && adduser -S nextjs -G nodejs

COPY --from=builder /app/next.config.mjs ./
COPY --from=builder /app/public ./public
COPY --from=builder --chown=nextjs:nodejs /app/.next/standalone ./
COPY --from=builder --chown=nextjs:nodejs /app/.next/static ./.next/static

USER nextjs
EXPOSE 3000
ENV PORT=3000
CMD ["node", "server.js"]""")

add_heading3(doc, "Stage 1 – deps")
add_body(doc,
    "Installs all npm dependencies. Using node:18-alpine (~5 MB vs ~900 MB for Debian) "
    "minimises the base size. Copying only package.json and package-lock.json first leverages "
    "Docker layer caching — if dependencies haven't changed, this expensive layer is reused "
    "without reinstalling."
)
add_heading3(doc, "Stage 2 – builder")
add_body(doc,
    "Copies installed node_modules and full source code, then runs npm run build (next build). "
    "The next.config.mjs specifies output: 'standalone', generating a .next/standalone "
    "directory — a self-contained Node.js server with only the files needed to run the app, "
    "excluding the full node_modules directory."
)
add_heading3(doc, "Stage 3 – runner")
add_body(doc,
    "The minimal production image. Security features: (1) A non-root user 'nextjs' is created "
    "and used to run the process, following the principle of least privilege. (2) Only the "
    "standalone server, public assets, and static files are copied — the full source code and "
    "node_modules are NOT in the final image. (3) The container listens on port 3000 and is "
    "started via node server.js (the standalone Next.js server)."
)

add_heading2(doc, "9.2  Docker Hub Repository Details")
add_spacer(doc, 2)
add_table(doc,
    headers=["Attribute", "Details"],
    rows=[
        ["Docker Hub Username",    "suicide768"],
        ["Repository Name",        "suicide768/nimmayatri-app"],
        ["Visibility",             "Public"],
        ["Image Tags",             "Build number (:1, :2, ...) + :latest always updated"],
        ["Base Image",             "node:18-alpine"],
        ["Auth Method",            "Jenkins Credentials (ID: dockerhub-creds) via withCredentials"],
        ["Final Image Digest",     "sha256:ce8053b55a74bb3f85c6cfbdcf9045bbb126e979a4c696162fe9aa637964cc5c"],
    ],
    col_widths=[2.0, 4.0]
)

add_heading2(doc, "9.3  Build and Push Commands")
add_code_block(doc,
"""# Build with build-number tag AND latest tag
docker build -t suicide768/nimmayatri-app:2 -t suicide768/nimmayatri-app:latest .

# Non-interactive Docker Hub login (password via stdin — never exposed in logs)
echo %DOCKER_PASS% | docker login -u %DOCKER_USER% --password-stdin

# Push both tags to Docker Hub
docker push suicide768/nimmayatri-app:2
docker push suicide768/nimmayatri-app:latest""")

add_heading2(doc, "9.4  Image Tags and Versioning")
add_bullet(doc, "Build number tag (:1, :2, :3 ...): Immutable, monotonically increasing. Enables rollback to any specific build (docker pull suicide768/nimmayatri-app:1).")
add_bullet(doc, ":latest tag: Always points to the most recently pushed image. Used as the default pull target when no tag is specified.")

add_heading2(doc, "9.5  Docker Hub Screenshots")
add_image(doc, pic("Screenshot 2026-05-18 185109.png"),
    "Fig 9.1 – VS Code terminal: Docker multi-stage build (18 build steps completed); "
    "docker run -p 3000:3000 nimmayatri-app confirms the container starts successfully on port 3000")
add_image(doc, pic("Screenshot 2026-05-18 191813.png"),
    "Fig 9.2 – Terminal: docker push suicide768/nimmayatri-app — all image layers pushed; "
    "final digest sha256:ce8053b55a74... confirms successful publication to Docker Hub")
add_image(doc, pic("Screenshot 2026-05-18 191905.png"),
    "Fig 9.3 – Docker Hub web interface: suicide768/nimmayatri-app repository visible, "
    "last pushed 1 minute ago, Visibility: Public")
add_image(doc, pic("Screenshot 2026-05-18 192316.png"),
    "Fig 9.4 – Terminal: docker pull suicide768/nimmayatri-app — image pulled from Docker Hub; "
    "'Status: Image is up to date' with digest verification confirms accessibility")

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# CH 10: DEPLOYMENT TO CLOUD PLATFORM
# ═══════════════════════════════════════════════════════
add_heading1(doc, "Chapter 10 – Deployment to Cloud Platform")
add_spacer(doc, 4)

add_heading2(doc, "10.1  Platform – Vercel")
add_body(doc,
    "Vercel is the cloud deployment platform used for the Nimma Yatri application. As the "
    "creators of Next.js, Vercel provides first-class Next.js support with zero-configuration "
    "setup, automatic HTTPS, global CDN, and seamless GitHub integration. The application is "
    "deployed in the Mumbai (bom1) region to minimise latency for Indian users."
)

add_heading2(doc, "10.2  Deployment Steps")
add_bullet(doc, "Step 1: Connect GitHub repository to Vercel by authorising the Vercel GitHub app on the harshendram account.")
add_bullet(doc, "Step 2: Import the Devops_Lab_Assign repository into Vercel's dashboard.")
add_bullet(doc, "Step 3: Vercel auto-detects the Next.js framework and configures build settings (npm run build, output: .next).")
add_bullet(doc, "Step 4: Add environment variables in Vercel project settings: NEXT_PUBLIC_GOOGLE_MAPS_API_KEY, NEXT_PUBLIC_GEMINI_API_KEY, GEMINI_API_KEY.")
add_bullet(doc, "Step 5: Click Deploy. Vercel builds and deploys to the global CDN with automatic HTTPS certificate.")
add_bullet(doc, "Step 6: Every subsequent push to the main branch triggers automatic Vercel re-deployment.")

add_heading2(doc, "10.3  Environment Configuration")
add_spacer(doc, 2)
add_table(doc,
    headers=["Environment Variable", "Purpose"],
    rows=[
        ["NEXT_PUBLIC_GOOGLE_MAPS_API_KEY", "Google Maps Platform — Places, Distance Matrix, Geocoding APIs (client-side)."],
        ["NEXT_PUBLIC_GEMINI_API_KEY",       "Gemini API for client-side features (Gemini Live WebSocket)."],
        ["GEMINI_API_KEY",                   "Gemini API for server-side API routes (chatbot endpoint)."],
    ],
    col_widths=[2.6, 3.4]
)

add_heading2(doc, "10.4  Deployment Details")
add_spacer(doc, 2)
add_table(doc,
    headers=["Attribute", "Value"],
    rows=[
        ["Platform",              "Vercel"],
        ["Production URL",        "https://nimmayatri.vercel.app"],
        ["Region",                "Mumbai, India (bom1)"],
        ["Framework",             "Next.js 14 (auto-detected)"],
        ["Build Command",         "npm run build"],
        ["HTTPS",                 "Automatic (Let's Encrypt)"],
        ["CDN",                   "Vercel Edge Network (Global)"],
        ["Lighthouse Score",      "94 / 100"],
        ["First Load JS",         "299 KB  (vendor chunk: 270 KB)"],
        ["Total Static Assets",   "141.77 MB  (21 files: 2 HTML, 13 JS, 1 CSS, 4 Images)"],
        ["Build Duration",        "52 seconds"],
        ["Total Deployments",     "6"],
        ["Trigger",               "GitHub push to main branch"],
    ],
    col_widths=[2.0, 4.0]
)

add_heading2(doc, "10.5  Deployment Screenshots")
add_image(doc, pic("Screenshot 2026-05-22 120330.png"),
    "Fig 10.1 – Vercel dashboard for nimmayatri project: Deployment Status = Ready (green), "
    "Production Domain = nimmayatri.vercel.app, created May 19 2026 by harshendram from "
    "commit cb350df; Firewall: Active, 0 block events; Observability: 0% Error Rate")
add_image(doc, pic("Screenshot 2026-05-22 120440.png"),
    "Fig 10.2 – Vercel build logs: Next.js 14.2.21 production compilation output — routing table "
    "(/ = 21 kB, /_not-found = 184 B, 4 serverless API functions), First Load JS = 299 kB, "
    "build time 52 seconds, 141.77 MB assets deployed")

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# CH 11: AUTOMATION USING WEBHOOKS
# ═══════════════════════════════════════════════════════
add_heading1(doc, "Chapter 11 – Automation using Webhooks")
add_spacer(doc, 4)

add_heading2(doc, "11.1  Overview")
add_body(doc,
    "Webhooks enable GitHub to actively push HTTP notifications to Jenkins the instant a push "
    "event occurs, rather than Jenkins polling GitHub at intervals. This results in near-"
    "instantaneous pipeline triggering — typically within 2–5 seconds of a git push. The "
    "ngrok tunnel bridges GitHub's internet-accessible webhook endpoint to the local Jenkins "
    "server running on a Windows machine without a public IP."
)

add_heading2(doc, "11.2  End-to-End Automation Flow")
add_spacer(doc, 2)
add_table(doc,
    headers=["Step", "Action", "Component"],
    rows=[
        ["1", "Developer runs: git push origin main", "Git / Developer Machine"],
        ["2", "GitHub receives the push and identifies all webhooks configured for the repository", "GitHub"],
        ["3", "GitHub sends HTTP POST to the Payload URL with a JSON body containing commit info, author, branch, and repository", "GitHub → ngrok"],
        ["4", "ngrok receives the HTTPS request and forwards it as HTTP to localhost:8080/github-webhook/", "ngrok Tunnel"],
        ["5", "Jenkins GitHub Integration Plugin validates the webhook payload and identifies matching pipeline jobs", "Jenkins"],
        ["6", "Jenkins triggers nimmayatri-pipeline (and nimmayatri-pipeline-no-push if configured)", "Jenkins Pipeline"],
        ["7", "All 9 pipeline stages execute sequentially from Checkout to Docker Push", "Jenkins"],
        ["8", "Trivy reports archived; Docker image pushed to Docker Hub; build status set to SUCCESS", "Jenkins / Docker Hub"],
    ],
    col_widths=[0.4, 3.4, 2.2]
)

add_heading2(doc, "11.3  ngrok Configuration Details")
add_spacer(doc, 2)
add_table(doc,
    headers=["Attribute", "Value"],
    rows=[
        ["ngrok Account",       "lordhshiva@gmail.com (Free Plan)"],
        ["Forwarding URL",      "https://semiobliviously-unevaporative-neriah.ngrok-free.dev"],
        ["Local Target",        "http://127.0.0.1:8080"],
        ["Region",              "India (in)"],
        ["Latency",             "16 ms"],
        ["Webhook Endpoint",    "https://<ngrok-url>/github-webhook/"],
        ["HTTP Response",       "200 OK (confirms Jenkins acknowledged delivery)"],
    ],
    col_widths=[2.0, 4.0]
)

add_heading2(doc, "11.4  Webhook Screenshots")
add_image(doc, pic("Screenshot 2026-05-18 210619.png"),
    "Fig 11.1 – GitHub Settings → Webhooks → Manage Webhook: Payload URL = ngrok HTTPS URL, "
    "Content type = application/json, SSL verification enabled, trigger = Just the push event")
add_image(doc, pic("Screenshot 2026-05-18 212826.png"),
    "Fig 11.2 – ngrok session: Forwarding URL mapped to http://127.0.0.1:8080, Region: India, "
    "HTTP request log showing POST /github-webhook/ → 200 OK — webhook successfully delivered")
add_image(doc, pic("Screenshot 2026-05-18 212056.png"),
    "Fig 11.3 – Jenkins Build #2 status: 'Started by GitHub push by harshendram' — definitive "
    "proof of automatic webhook-triggered pipeline execution (vs manual trigger which shows 'Started by user')")
add_image(doc, pic("Screenshot 2026-05-18 212021.png"),
    "Fig 11.4 – Jenkins Console Output: 'Started by GitHub push by harshendram'; "
    "Jenkinsfile obtained from git at the triggering commit's HEAD, pipeline stages beginning")

add_heading2(doc, "11.5  Verification of Automatic Execution")
add_body(doc, "Automatic webhook-based execution was verified through two independent sources:")
add_bullet(doc, "Jenkins build metadata: The build status page for Build #2 shows 'Started by GitHub push by harshendram'. This text is set exclusively by the Jenkins GitHub Integration Plugin upon receipt of a valid webhook payload — it cannot appear for manual builds.")
add_bullet(doc, "ngrok request log: The ngrok dashboard at http://127.0.0.1:4040 records POST /github-webhook/ with HTTP 200 response, confirming GitHub delivered the payload and Jenkins acknowledged receipt.")

add_page_break(doc)

# ═══════════════════════════════════════════════════════
# CH 12: CONCLUSION
# ═══════════════════════════════════════════════════════
add_heading1(doc, "Chapter 12 – Conclusion")
add_spacer(doc, 4)

add_heading2(doc, "12.1  Summary")
add_body(doc,
    "This project successfully designed and implemented a complete, production-quality DevOps "
    "CI/CD pipeline for the Nimma Yatri Next.js application. The pipeline integrates GitHub "
    "source control, Jenkins automation, ESLint code quality enforcement, Trivy security "
    "scanning, Docker multi-stage containerisation, Docker Hub image registry, Vercel cloud "
    "deployment, and GitHub webhook-based automatic triggering into a cohesive, fully "
    "automated software delivery system."
)
add_body(doc,
    "Both pipeline variants — the full pipeline with Docker push (Build #2: 3 min 54 sec) "
    "and the no-push testing variant (Build #2: 1 min 19 sec) — completed successfully with "
    "green status. The Docker image suicide768/nimmayatri-app:latest was published to Docker Hub, "
    "both Trivy reports were archived as Jenkins artefacts, and the webhook-based automation "
    "was confirmed by the 'Started by GitHub push by harshendram' build trigger message."
)

add_heading2(doc, "12.2  Lessons Learnt")
add_bullet(doc, "Shift quality left: Placing ESLint checks and Trivy scans early in the pipeline ensures issues are caught before they can reach production, minimising the cost of defect detection.")
add_bullet(doc, "Security scanning reveals hidden vulnerabilities: The Trivy scan uncovered 28 CVEs in package-lock.json and ~58 in the Docker image, including a CRITICAL Next.js authorization bypass (CVE-2025-29927) that was not apparent from the application's functionality alone.")
add_bullet(doc, "Multi-stage Dockerfiles are essential for production: The three-stage build produced a significantly smaller and more secure image by excluding source code, build tools, and full node_modules from the runner stage, running as a non-root user.")
add_bullet(doc, "Local Jenkins requires bridging infrastructure: Running Jenkins locally without a public IP required ngrok for webhook delivery. In a production DevOps setup, Jenkins would run on a cloud VM (AWS EC2, GCP, Azure) with a static IP, eliminating this dependency.")
add_bullet(doc, "Jenkins Credentials store is non-negotiable: Storing Docker Hub credentials in Jenkins and injecting them via withCredentials prevents credentials from appearing in console logs, build history, or source code.")
add_bullet(doc, "Pipeline as Code (Jenkinsfile in SCM) is the correct model: Version-controlling the Jenkinsfile alongside application code ensures pipeline changes are auditable, peer-reviewable, and always in sync with the code they deploy.")
add_bullet(doc, "Two pipeline variants serve different purposes: The no-push variant proved valuable for rapid development iteration, avoiding unnecessary Docker Hub image pollution during testing.")
add_bullet(doc, "DevOps tools require investment but deliver compounding returns: Initial configuration of Jenkins, Trivy, Docker multi-stage builds, and webhook integration required significant time. However, once configured, the system runs reliably and automatically for every subsequent commit with zero developer overhead.")

add_heading2(doc, "12.3  Future Enhancements")
add_bullet(doc, "Upgrade Next.js to ≥ 14.2.25 to resolve the CRITICAL CVE-2025-29927 authorization bypass vulnerability.")
add_bullet(doc, "Add Jest unit and integration tests with a dedicated Test stage in the pipeline.")
add_bullet(doc, "Integrate SonarQube Cloud for deeper code quality metrics (coverage, duplication, code smells).")
add_bullet(doc, "Migrate Jenkins to a cloud VM to eliminate ngrok dependency and enable 24/7 persistent pipeline availability.")
add_bullet(doc, "Add automated deployment from Docker Hub to Render or Railway as an additional cloud target.")
add_bullet(doc, "Implement Docker Compose for local multi-container development environment management.")

# ─────────────────────────────────────────────
# SAVE DOCX
# ─────────────────────────────────────────────
doc.save(DOCX)
print("DOCX saved: " + DOCX)

# ─────────────────────────────────────────────
# EXPORT PDF via win32com (updates TOC field first)
# ─────────────────────────────────────────────
print("Opening in Word to update TOC and export PDF...")
try:
    import win32com.client as win32
    word = win32.Dispatch('Word.Application')
    word.Visible = False
    wd_doc = word.Documents.Open(DOCX)
    wd_doc.TablesOfContents(1).Update()
    wd_doc.Save()
    wd_doc.ExportAsFixedFormat(
        OutputFileName   = PDF,
        ExportFormat     = 17,       # wdExportFormatPDF
        OpenAfterExport  = False,
        OptimizeFor      = 0,        # wdExportOptimizeForPrint
        Range            = 0,        # wdExportAllDocument
        IncludeDocProps  = True,
        CreateBookmarks  = 1,        # wdExportCreateHeadingBookmarks
        DocStructureTags = True,
        BitmapMissingFonts = True,
        UseISO19005_1    = False,
    )
    wd_doc.Close(False)
    word.Quit()
    print("PDF saved: " + PDF)
except Exception as e:
    print("win32com PDF export failed: " + str(e))
    print("Falling back to docx2pdf...")
    from docx2pdf import convert
    convert(DOCX, PDF)
    print("PDF saved: " + PDF)
